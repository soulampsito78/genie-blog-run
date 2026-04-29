#!/usr/bin/env python3
"""
Owner image contract (today_genie):

- TOP slot: fixed reference JPEG only (GENIE_REF_today_genie_master_v1.jpg). Not generated.
- BOTTOM slot: TPO-responsive image from fresh full-feed run (outdoor prompt + mood prefix),
  same-person continuity via the same reference file passed into generate_image_file.

Writes versioned GENIE_EMAIL_today_genie_bottom_v{N}.jpg, ops/preview contract + evidence,
owner_review HTML (inside white card: **title h1 → fixed top → body → generated bottom → footer**;
mode label text is only in the gray intro line outside the card) and DOCX (title → [IMAGE_TOP] → body → [IMAGE_BOTTOM]).

CLI:
  (default)   full-feed text run + bottom image + contract + owner artifacts
  --artifacts-only   bottom image + artifacts from saved passing contract (no text model)
  --rebuild-preview-only   re-read contract JSON; rewrite HTML + DOCX + normalization evidence only
"""
from __future__ import annotations

import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from docx.shared import Pt
from fastapi import HTTPException

_REPO = Path(__file__).resolve().parents[1]
# Locked top slot: same default as TodayGenieImageJobRequest.reference_image_path in main.py.
FIXED_TOP_EMAIL_FILENAME = "GENIE_REF_today_genie_master_v1.jpg"
FIXED_TOP_ALT = (
    "GENIE today_genie — fixed master reference for email top slot (not generated; locked asset)."
)

OWNER_PREVIEW_HTML = _REPO / "ops/preview/owner_review_today_genie_email_preview.html"
OWNER_BODY_DOCX = _REPO / "ops/preview/owner_review_today_genie_body.docx"
CONTRACT_JSON = _REPO / "ops/preview/today_genie_contract_run_latest.json"
PIPELINE_NORMALIZATION_EVIDENCE = _REPO / "ops/preview/owner_review_pipeline_normalization_evidence.json"
# Preview HTML lives under ops/preview/ — img src must resolve under file:// from that directory.
OWNER_IMG_SRC_BASE = "../../static/email"

_FEED_FILES = [
    ("TODAY_GENIE_OVERNIGHT_US_MARKET_JSON", "overnight_us_market.json"),
    ("TODAY_GENIE_MACRO_INDICATORS_JSON", "macro_indicators.json"),
    ("TODAY_GENIE_TOP_MACRO_ISSUES_JSON", "top_macro_issues.json"),
    ("TODAY_GENIE_TOP_MARKET_NEWS_JSON", "top_market_news.json"),
    ("TODAY_GENIE_KOREA_MARKET_SCHEDULE_JSON", "korea_market_schedule.json"),
    ("TODAY_GENIE_RISK_FACTORS_JSON", "risk_factors.json"),
    ("TODAY_GENIE_KOREA_JAPAN_INDICES_JSON", "korea_japan_indices.json"),
]


def _load_feeds_or_exit() -> None:
    feeds_dir = _REPO / "ops/feeds"
    for env_key, fname in _FEED_FILES:
        path = feeds_dir / fname
        if not path.is_file():
            print(json.dumps({"error": "missing_feed_file", "path": str(path)}, ensure_ascii=False))
            sys.exit(10)
        val = json.loads(path.read_text(encoding="utf-8"))
        if val is None or (isinstance(val, dict) and len(val) == 0) or (isinstance(val, list) and len(val) == 0):
            print(json.dumps({"error": "empty_feed", "path": str(path)}, ensure_ascii=False))
            sys.exit(12)
        os.environ[env_key] = json.dumps(val, ensure_ascii=False, separators=(",", ":"))


def _next_versioned_bottom_path(email_dir: Path) -> Path:
    """GENIE_EMAIL_today_genie_bottom_v{N}.jpg with N one above max existing vN."""
    pat = re.compile(r"GENIE_EMAIL_today_genie_bottom_v(\d+)\.jpg$", re.IGNORECASE)
    max_v = 0
    for p in email_dir.glob("GENIE_EMAIL_today_genie_bottom_v*.jpg"):
        m = pat.match(p.name)
        if m:
            max_v = max(max_v, int(m.group(1)))
    return email_dir / f"GENIE_EMAIL_today_genie_bottom_v{max_v + 1}.jpg"


def _write_owner_html(data: dict, top_name: str, bottom_name: str, out_path: Path) -> None:
    from renderers import (
        TODAY_EMAIL_CLOSING_CRITERION,
        _build_today_genie_email_editorial_html,
        _safe,
    )

    editorial = _build_today_genie_email_editorial_html(data)
    m = re.match(r"^(.+?)\s*(<h1[^>]*>.*?</h1>)", editorial.strip(), re.DOTALL)
    if not m:
        raise RuntimeError("Could not parse editorial header/title")
    header_label_html, h1_html = m.group(1).strip(), m.group(2).strip()
    rest_after_h1 = editorial[m.end() :].lstrip()
    # Strip tags for intro-only mode line (keeps owner context without placing label before title in card).
    mode_plain = re.sub(r"<[^>]+>", "", header_label_html).strip() or "[장전 브리핑]"
    criterion_p = (
        f'<p style="margin:0 0 14px 0;font-size:12px;line-height:1.6;color:#666;">'
        f"{_safe(TODAY_EMAIL_CLOSING_CRITERION)}</p>"
    )
    idx = rest_after_h1.find(criterion_p)
    if idx == -1:
        main_body = rest_after_h1
        footer_html = ""
    else:
        main_body = rest_after_h1[:idx]
        footer_html = rest_after_h1[idx:]

    alt_top = (
        _safe(FIXED_TOP_ALT)
        if top_name == FIXED_TOP_EMAIL_FILENAME
        else _safe(data.get("image_alt_top") or "today_genie email top")
    )
    alt_bot = _safe(data.get("image_alt_bottom") or "today_genie email bottom")
    top_img = (
        f'<div style="margin:0 0 20px 0;text-align:center;">'
        f'<img src="{OWNER_IMG_SRC_BASE}/{top_name}" alt="{alt_top}" '
        'style="max-width:100%;height:auto;display:block;margin:0 auto;"/>'
        "</div>"
    )
    bottom_img = (
        f'<div style="margin:20px 0 0 0;text-align:center;">'
        f'<img src="{OWNER_IMG_SRC_BASE}/{bottom_name}" alt="{alt_bot}" '
        'style="max-width:100%;height:auto;display:block;margin:0 auto;"/>'
        "</div>"
    )
    inner = (
        f'<div style="margin:0;padding:0;background:#ffffff;">'
        f'<div style="max-width:640px;width:100%;margin:0 auto;background:#ffffff;'
        f"font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Arial,sans-serif;"
        f'font-size:15px;line-height:1.75;color:#1a1a1a;">'
        f'<div style="padding:24px;">'
        f"{h1_html}\n{top_img}\n{main_body}\n{bottom_img}\n{footer_html}"
        f"</div></div></div>"
    )
    html_out = f"""<!DOCTYPE html>
<html lang="ko">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>Owner Review — today_genie email preview</title>
</head>
<body style="margin:0;padding:24px;background:#e2e8f0;">
  <p style="margin:0 0 16px 0;font-size:13px;color:#475569;font-family:system-ui,sans-serif;">{mode_plain} — Local owner-review pack. Card order: <strong>title</strong> → fixed TOP <code>{top_name}</code> → body → generated BOTTOM <code>{bottom_name}</code> → footer. Images use <code>src="{OWNER_IMG_SRC_BASE}/…</code> relative to this file so <code>file://</code> preview resolves.</p>
  <div style="max-width:640px;margin:0 auto;background:#fff;padding:24px;font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Arial,sans-serif;">
{inner}
  </div>
</body>
</html>
"""
    out_path.write_text(html_out, encoding="utf-8")


def _write_owner_docx(data: dict, out_path: Path) -> None:
    from renderers import TODAY_EMAIL_CLOSING_CRITERION

    title = (data.get("title") or "").strip()
    opening = (data.get("opening_summary") or data.get("summary") or "").strip()
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)
    doc.add_heading(title, level=1)
    doc.add_paragraph("[IMAGE_TOP]")
    if opening:
        doc.add_paragraph(opening)

    doc.add_heading("TOP 3 뉴스", level=2)
    for it in data.get("top_3_news") or []:
        if not isinstance(it, dict):
            continue
        doc.add_paragraph(it.get("headline") or "", style="List Bullet")
        if it.get("what_happened"):
            doc.add_paragraph(it["what_happened"], style="List Bullet 2")
        if it.get("key_numbers"):
            doc.add_paragraph(f"핵심 수치: {it['key_numbers']}", style="List Bullet 2")
        if it.get("why_it_matters_today"):
            doc.add_paragraph(f"오늘 의미: {it['why_it_matters_today']}", style="List Bullet 2")

    doc.add_heading("핵심 숫자 5", level=2)
    for it in data.get("key_numbers_5") or []:
        if not isinstance(it, dict):
            continue
        doc.add_paragraph(it.get("headline") or "", style="List Bullet")
        if it.get("number"):
            doc.add_paragraph(f"수치: {it['number']}", style="List Bullet 2")
        if it.get("why_it_matters"):
            doc.add_paragraph(it["why_it_matters"], style="List Bullet 2")

    for label, key in (
        ("오늘 먼저 볼 것 3", "today_watch_first_3"),
        ("과해석 금지 3", "do_not_overread_3"),
    ):
        items = data.get(key) or []
        if items:
            doc.add_heading(label, level=2)
            for it in items:
                if isinstance(it, str):
                    doc.add_paragraph(it, style="List Bullet")

    closing = (data.get("final_decision_line") or data.get("closing_message") or "").strip()
    if closing:
        doc.add_heading("오늘의 결정 기준", level=2)
        doc.add_paragraph(closing)

    doc.add_paragraph("[IMAGE_BOTTOM]")
    doc.add_paragraph(TODAY_EMAIL_CLOSING_CRITERION)
    doc.add_paragraph("© Heemang & Tobak. All rights reserved.")
    for line in data.get("disclaimer_block") or []:
        if isinstance(line, str) and line.strip():
            doc.add_paragraph(line.strip())
    tags = data.get("hashtags") or []
    if tags:
        doc.add_heading("해시태그", level=2)
        parts = []
        for t in tags:
            s = str(t).strip()
            if not s:
                continue
            parts.append(s if s.startswith("#") else f"#{s}")
        doc.add_paragraph(" ".join(parts))
    doc.save(out_path)


def _emit_owner_pipeline_normalization_evidence() -> None:
    """Single compact audit record for the local owner-review pipeline (today_genie)."""
    contract: dict = {}
    if CONTRACT_JSON.is_file():
        contract = json.loads(CONTRACT_JSON.read_text(encoding="utf-8"))
    ic = contract.get("image_contract") or {}
    top_repo = (ic.get("top_slot") or {}).get("path") or f"static/email/{FIXED_TOP_EMAIL_FILENAME}"
    bottom_repo = (ic.get("bottom_slot") or {}).get("path") or ""
    top_abs = (_REPO / top_repo.lstrip("/")).resolve() if top_repo else None
    bottom_abs = (_REPO / bottom_repo.lstrip("/")).resolve() if bottom_repo else None

    html_text = OWNER_PREVIEW_HTML.read_text(encoding="utf-8") if OWNER_PREVIEW_HTML.is_file() else ""
    img_srcs = re.findall(r'<img[^>]+src="([^"]+)"', html_text)
    html_img_audit: list[dict[str, object]] = []
    for src in img_srcs:
        resolved = (OWNER_PREVIEW_HTML.parent / src).resolve()
        html_img_audit.append(
            {
                "src": src,
                "resolved_path": str(resolved),
                "exists": resolved.is_file(),
                "bytes": resolved.stat().st_size if resolved.is_file() else None,
            }
        )

    docx_top_ok = False
    docx_bottom_ok = False
    docx_bottom_indices: list[int] = []
    embedded_images = -1
    if OWNER_BODY_DOCX.is_file():
        doc = Document(str(OWNER_BODY_DOCX))
        embedded_images = len(doc.inline_shapes)
        paras = doc.paragraphs
        if len(paras) > 1:
            docx_top_ok = str(paras[0].style.name).startswith("Heading") and paras[1].text.strip() == "[IMAGE_TOP]"
        docx_bottom_indices = [i for i, p in enumerate(paras) if p.text.strip() == "[IMAGE_BOTTOM]"]
        docx_bottom_ok = len(docx_bottom_indices) == 1 and docx_bottom_indices[0] > 2

    top_visible = bool(html_img_audit and html_img_audit[0].get("exists"))
    bottom_visible = bool(len(html_img_audit) > 1 and html_img_audit[1].get("exists"))

    payload = {
        "giant_step": "today_genie_local_owner_review_pipeline_normalization",
        "completed": True,
        "timestamp_kst": datetime.now(ZoneInfo("Asia/Seoul")).isoformat(),
        "top_image_contract": {
            "repo_path": top_repo,
            "absolute_path": str(top_abs) if top_abs else None,
            "kind": "fixed_reference",
            "generated": False,
            "note": "Top email slot uses locked master reference JPEG; never Vertex-generated for owner pack.",
        },
        "bottom_image_contract": {
            "repo_path": bottom_repo,
            "absolute_path": str(bottom_abs) if bottom_abs else None,
            "kind": "generated_tpo_bottom",
            "generated": True,
            "source_json": str(CONTRACT_JSON.relative_to(_REPO)),
            "reference_for_continuity": (ic.get("bottom_slot") or {}).get("reference_for_continuity"),
        },
        "html_binding": {
            "preview_path": str(OWNER_PREVIEW_HTML.relative_to(_REPO)),
            "img_src_values": img_srcs,
            "resolved_from_preview_dir": html_img_audit,
            "top_visible_file_url": top_visible,
            "bottom_visible_file_url": bottom_visible,
        },
        "docx": {
            "path": str(OWNER_BODY_DOCX.relative_to(_REPO)),
            "image_top_below_title": docx_top_ok,
            "image_bottom_lower_slot": docx_bottom_ok,
            "image_bottom_paragraph_indices": docx_bottom_indices,
            "embedded_image_shapes": embedded_images,
        },
    }
    PIPELINE_NORMALIZATION_EVIDENCE.parent.mkdir(parents=True, exist_ok=True)
    PIPELINE_NORMALIZATION_EVIDENCE.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def rebuild_preview_only_main() -> int:
    """Re-read passing contract; rewrite owner HTML + DOCX + normalization evidence (no APIs)."""
    os.chdir(_REPO)
    sys.path.insert(0, str(_REPO))
    if not CONTRACT_JSON.is_file():
        print(json.dumps({"error": "missing_contract", "path": str(CONTRACT_JSON)}, ensure_ascii=False))
        return 80
    raw = json.loads(CONTRACT_JSON.read_text(encoding="utf-8"))
    if raw.get("validation_result") != "pass":
        print(
            json.dumps(
                {"error": "contract_not_pass", "validation_result": raw.get("validation_result")},
                ensure_ascii=False,
            )
        )
        return 81
    data = raw.get("data")
    if not isinstance(data, dict):
        print(json.dumps({"error": "contract_missing_data"}, ensure_ascii=False))
        return 82
    ic = raw.get("image_contract") or {}
    top_fn = Path((ic.get("top_slot") or {}).get("path") or f"static/email/{FIXED_TOP_EMAIL_FILENAME}").name
    bpath = (ic.get("bottom_slot") or {}).get("path") or ""
    bottom_fn = Path(bpath).name if bpath else ""
    if not bottom_fn:
        print(json.dumps({"error": "image_contract_missing_bottom"}, ensure_ascii=False))
        return 83

    meta = dict(raw.get("run_meta") or {})
    meta["owner_pipeline_normalization_kst"] = datetime.now(ZoneInfo("Asia/Seoul")).isoformat()
    meta["owner_image_contract_audit"] = {
        "top_slot_statement": "static/email/GENIE_REF_today_genie_master_v1.jpg is the fixed reference; not Vertex-generated for the email top slot.",
        "bottom_slot_statement": f"{bpath or f'static/email/{bottom_fn}'} is the TPO-generated bottom tied to this contract JSON.",
        "supersedes": "Any prior owner-pack step that treated GENIE_EMAIL_today_genie_top_v2.jpg as the locked top or as non-reference output.",
    }
    raw["run_meta"] = meta
    CONTRACT_JSON.write_text(json.dumps(raw, ensure_ascii=False, indent=2), encoding="utf-8")

    OWNER_PREVIEW_HTML.parent.mkdir(parents=True, exist_ok=True)
    _write_owner_html(data, top_fn, bottom_fn, OWNER_PREVIEW_HTML)
    _write_owner_docx(data, OWNER_BODY_DOCX)
    _emit_owner_pipeline_normalization_evidence()
    print(
        json.dumps(
            {
                "ok": True,
                "rebuild_preview_only": True,
                "preview_html": str(OWNER_PREVIEW_HTML.resolve()),
                "docx": str(OWNER_BODY_DOCX.resolve()),
                "evidence": str(PIPELINE_NORMALIZATION_EVIDENCE.resolve()),
            },
            ensure_ascii=False,
        )
    )
    return 0


def artifacts_only_main() -> int:
    """
    Regenerate bottom image + owner HTML/DOCX from the last saved passing contract on disk.
    Does not call the text model. Use when a live full run cannot reach validation pass but
    ops/preview/today_genie_contract_run_latest.json already holds a pass snapshot.
    """
    os.chdir(_REPO)
    sys.path.insert(0, str(_REPO))

    from image_exec_suffixes import today_genie_suffix_outdoor_daily
    from image_generator import generate_image_file
    from main import (
        PROJECT_ID,
        VERTEX_IMAGE_MODEL,
        VERTEX_LOCATION,
        _mood_prefix_for_image_prompts,
    )
    from renderers import TODAY_EMAIL_CLOSING_CRITERION

    if not CONTRACT_JSON.is_file():
        print(json.dumps({"error": "no_saved_contract", "path": str(CONTRACT_JSON)}, ensure_ascii=False))
        return 80
    raw = json.loads(CONTRACT_JSON.read_text(encoding="utf-8"))
    if raw.get("validation_result") != "pass":
        print(
            json.dumps(
                {
                    "error": "saved_contract_not_pass",
                    "validation_result": raw.get("validation_result"),
                },
                ensure_ascii=False,
            )
        )
        return 81
    data = raw.get("data")
    if not isinstance(data, dict):
        print(json.dumps({"error": "saved_contract_missing_data"}, ensure_ascii=False))
        return 82

    kst = datetime.now(ZoneInfo("Asia/Seoul"))
    run_id = kst.isoformat()
    mood_prefix = _mood_prefix_for_image_prompts(data)
    p_out = (data.get("image_prompt_outdoor") or "").strip()
    if not p_out:
        print(json.dumps({"error": "missing_image_prompt_outdoor"}, ensure_ascii=False))
        return 3

    email_dir = _REPO / "static" / "email"
    email_dir.mkdir(parents=True, exist_ok=True)
    fixed_top_path = email_dir / FIXED_TOP_EMAIL_FILENAME
    bottom_out = _next_versioned_bottom_path(email_dir)
    if not fixed_top_path.is_file():
        print(json.dumps({"error": "fixed_top_missing", "path": str(fixed_top_path)}, ensure_ascii=False))
        return 5

    ri_img = raw.get("runtime_input") if isinstance(raw.get("runtime_input"), dict) else {}
    try:
        generate_image_file(
            prompt=mood_prefix
            + p_out
            + "\n\n"
            + today_genie_suffix_outdoor_daily(ri_img, variation_seed=run_id.replace(":", "-")),
            output_path=bottom_out,
            model_name=VERTEX_IMAGE_MODEL,
            reference_image_path=fixed_top_path,
            project_id=PROJECT_ID or None,
            location=VERTEX_LOCATION,
        )
        try:
            from genie_image_overlay import apply_today_genie_brand_footer

            apply_today_genie_brand_footer(bottom_out)
        except Exception:
            pass
    except Exception as e:
        print(json.dumps({"error": "image_generation_failed", "message": f"{type(e).__name__}: {e}"}, ensure_ascii=False))
        return 4

    image_contract = {
        "top_slot": {
            "kind": "fixed_reference",
            "path": f"static/email/{FIXED_TOP_EMAIL_FILENAME}",
            "generated": False,
        },
        "bottom_slot": {
            "kind": "generated_tpo_bottom",
            "path": f"static/email/{bottom_out.name}",
            "generated": True,
            "reference_for_continuity": f"static/email/{FIXED_TOP_EMAIL_FILENAME}",
        },
    }
    meta = dict(raw.get("run_meta") or {})
    meta["artifacts_only_refresh_kst"] = run_id
    meta["artifacts_only_note"] = (
        "Bottom slot regenerated with Vertex from saved passing contract data; text model not re-run."
    )
    raw["run_meta"] = meta
    raw["image_contract"] = image_contract

    OWNER_PREVIEW_HTML.parent.mkdir(parents=True, exist_ok=True)
    CONTRACT_JSON.write_text(json.dumps(raw, ensure_ascii=False, indent=2), encoding="utf-8")

    evidence_path = _REPO / "ops/preview/owner_review_image_contract_latest.json"
    evidence_path.write_text(
        json.dumps(
            {
                "run_meta": meta,
                "source": "artifacts_only_from_saved_contract",
                "validation_result": raw.get("validation_result"),
                "image_briefing_mood_state": data.get("image_briefing_mood_state"),
                "image_mood_basis": data.get("image_mood_basis"),
                "image_contract": image_contract,
                "criterion_line_in_html": TODAY_EMAIL_CLOSING_CRITERION[:40] + "…",
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    _write_owner_html(
        data,
        FIXED_TOP_EMAIL_FILENAME,
        bottom_out.name,
        OWNER_PREVIEW_HTML,
    )
    _write_owner_docx(data, OWNER_BODY_DOCX)
    _emit_owner_pipeline_normalization_evidence()

    print(
        json.dumps(
            {
                "ok": True,
                "artifacts_only": True,
                "contract": str(CONTRACT_JSON.resolve()),
                "evidence": str(evidence_path.resolve()),
                "top_fixed_reference": str(fixed_top_path.resolve()),
                "bottom_generated": str(bottom_out.resolve()),
            },
            ensure_ascii=False,
        )
    )
    return 0


def main() -> int:
    os.chdir(_REPO)
    sys.path.insert(0, str(_REPO))
    _load_feeds_or_exit()

    from image_exec_suffixes import today_genie_suffix_outdoor_daily
    from image_generator import generate_image_file
    from main import (
        PROJECT_ID,
        VERTEX_IMAGE_MODEL,
        VERTEX_LOCATION,
        _mood_prefix_for_image_prompts,
        build_runtime_input,
        enforce_today_genie_market_snapshot_from_feeds,
        response_issues,
        run_today_genie_text_pipeline,
        stabilize_today_genie_validation_fields,
        validate_today_genie,
    )
    from renderers import TODAY_EMAIL_CLOSING_CRITERION

    mode = "today_genie"
    kst = datetime.now(ZoneInfo("Asia/Seoul"))
    run_id = kst.isoformat()

    ri = build_runtime_input(mode)
    if ri.get("input_feed_status") != "full":
        print(
            json.dumps(
                {
                    "error": "feeds_not_full",
                    "input_feed_status": ri.get("input_feed_status"),
                },
                ensure_ascii=False,
            )
        )
        return 20

    max_attempts = 6
    data: dict = {}
    raw_text = ""
    validation = None
    attempts_log: list[dict] = []

    for attempt in range(1, max_attempts + 1):
        try:
            data, raw_text, _ = run_today_genie_text_pipeline(ri)
            data = enforce_today_genie_market_snapshot_from_feeds(data, ri)
            data = stabilize_today_genie_validation_fields(data, ri)
            validation = validate_today_genie(data, ri)
        except HTTPException as e:
            detail = e.detail if isinstance(e.detail, dict) else {"message": str(e.detail)}
            attempts_log.append(
                {
                    "attempt": attempt,
                    "result": "parse_error",
                    "message": detail.get("message", str(detail)),
                }
            )
            continue
        except Exception as e:
            attempts_log.append(
                {"attempt": attempt, "result": "runtime_error", "message": f"{type(e).__name__}: {e}"}
            )
            continue

        attempts_log.append(
            {
                "attempt": attempt,
                "result": validation.result,
                "codes": [i.code for i in validation.issues[:12]],
            }
        )
        if validation.result in ("pass", "draft_only"):
            break
    else:
        print(
            json.dumps(
                {
                    "error": "validation_retry_exhausted",
                    "attempts": attempts_log,
                },
                ensure_ascii=False,
                indent=2,
            )
        )
        return 2

    if validation is None or validation.result not in ("pass", "draft_only"):
        print(json.dumps({"error": "no_pass", "attempts": attempts_log}, ensure_ascii=False, indent=2))
        return 2

    mood_prefix = _mood_prefix_for_image_prompts(data)
    p_out = (data.get("image_prompt_outdoor") or "").strip()
    if not p_out:
        print(json.dumps({"error": "missing_image_prompt_outdoor"}, ensure_ascii=False))
        return 3

    email_dir = _REPO / "static" / "email"
    email_dir.mkdir(parents=True, exist_ok=True)
    fixed_top_path = email_dir / FIXED_TOP_EMAIL_FILENAME
    bottom_out = _next_versioned_bottom_path(email_dir)
    if not fixed_top_path.is_file():
        print(json.dumps({"error": "fixed_top_missing", "path": str(fixed_top_path)}, ensure_ascii=False))
        return 5

    try:
        generate_image_file(
            prompt=mood_prefix
            + p_out
            + "\n\n"
            + today_genie_suffix_outdoor_daily(dict(ri), variation_seed=run_id.replace(":", "-")),
            output_path=bottom_out,
            model_name=VERTEX_IMAGE_MODEL,
            reference_image_path=fixed_top_path,
            project_id=PROJECT_ID or None,
            location=VERTEX_LOCATION,
        )
        try:
            from genie_image_overlay import apply_today_genie_brand_footer

            apply_today_genie_brand_footer(bottom_out)
        except Exception:
            pass
    except Exception as e:
        print(json.dumps({"error": "image_generation_failed", "message": f"{type(e).__name__}: {e}"}, ensure_ascii=False))
        return 4

    contract = {
        "run_meta": {"datetime": run_id, "mode": mode, "script": "run_owner_review_full_tpo_v2.py"},
        "runtime_input": ri,
        "attempts": attempts_log,
        "validation_result": validation.result,
        "issues": response_issues(validation.issues),
        "data": data,
        "image_contract": {
            "top_slot": {
                "kind": "fixed_reference",
                "path": f"static/email/{FIXED_TOP_EMAIL_FILENAME}",
                "generated": False,
            },
            "bottom_slot": {
                "kind": "generated_tpo_bottom",
                "path": f"static/email/{bottom_out.name}",
                "generated": True,
                "reference_for_continuity": f"static/email/{FIXED_TOP_EMAIL_FILENAME}",
            },
        },
    }
    OWNER_PREVIEW_HTML.parent.mkdir(parents=True, exist_ok=True)
    CONTRACT_JSON.write_text(
        json.dumps(contract, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    evidence_path = _REPO / "ops/preview/owner_review_image_contract_latest.json"
    evidence_path.write_text(
        json.dumps(
            {
                "run_meta": contract["run_meta"],
                "input_feed_status": ri.get("input_feed_status"),
                "validation_result": validation.result,
                "image_briefing_mood_state": data.get("image_briefing_mood_state"),
                "image_mood_basis": data.get("image_mood_basis"),
                "image_contract": contract["image_contract"],
                "criterion_line_in_html": TODAY_EMAIL_CLOSING_CRITERION[:40] + "…",
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    _write_owner_html(
        data,
        FIXED_TOP_EMAIL_FILENAME,
        bottom_out.name,
        OWNER_PREVIEW_HTML,
    )
    _write_owner_docx(data, OWNER_BODY_DOCX)
    _emit_owner_pipeline_normalization_evidence()

    print(
        json.dumps(
            {
                "ok": True,
                "contract": str(CONTRACT_JSON.resolve()),
                "evidence": str(evidence_path.resolve()),
                "top_fixed_reference": str(fixed_top_path.resolve()),
                "bottom_generated": str(bottom_out.resolve()),
                "mood": data.get("image_briefing_mood_state"),
            },
            ensure_ascii=False,
        )
    )
    return 0


if __name__ == "__main__":
    if "--rebuild-preview-only" in sys.argv:
        raise SystemExit(rebuild_preview_only_main())
    if "--artifacts-only" in sys.argv:
        raise SystemExit(artifacts_only_main())
    raise SystemExit(main())
